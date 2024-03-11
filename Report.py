import customtkinter
import tkinter
import re
from db import *
from CTkMessagebox import CTkMessagebox
from hPyT import *
from PIL import Image
import os
from tkcalendar import Calendar
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from docxtpl import DocxTemplate
customtkinter.set_appearance_mode("Dark")
customtkinter.set_default_color_theme("blue")


sql = SQL(server='DDLAPTOP\SQLEXPRESS', database='PC')
sql.connect()
image_path = os.path.join(os.path.dirname(os.path.realpath(__file__)))


class Reports(customtkinter.CTkToplevel):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.create_widgets()
        
    
    def create_widgets(self):
        self.title("Edit Computer")
        self.geometry("850x  450")
        self.minsize  (850,  450)
        self.maxsize  (850,  450)
        maximize_minimize_button.hide(self)
        self.grid_columnconfigure((1),weight=1)
        self.grid_rowconfigure((0,1),weight=1)
        self.create_frame1()
        self.create_frame2()
        self.create_frame3()
        self.create_frame4()
            #отчет по все инфо на выбранный пк
            #
            #отчет по ремонтам за промежуток дат
        
        
        
        
        
    def create_frame1(self):
        self.frame1 = customtkinter.CTkFrame(self,height=320)
        self.frame1.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        self.frame1.grid_columnconfigure((0,1,2,3),weight=1)
        customtkinter.CTkLabel(master=self.frame1,text="Отчет по статусам работы",fg_color="gray30", font=("Arial", 14)).grid(row=0,columnspan=4, column=0, padx=10, pady=10, sticky="ew")

        
        customtkinter.CTkLabel(master=self.frame1,text="Выберите дату начала периода").grid(row=1, column=0, padx=10,columnspan=4, pady=10, sticky="ew")
        self.DataStart_entry = customtkinter.CTkEntry(master=self.frame1, placeholder_text="ГГГГ-ММ-ДД")
        self.DataStart_entry.grid(row=2, column=1,columnspan=2, padx=10, pady=10, sticky="ew")
        
        self.image_icon_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "calendarICO.png")), size=(20, 20))
        self.CalendarOpenButton = customtkinter.CTkButton(master=self.frame1, text="Выбрать дату",image=self.image_icon_image, command=lambda: [self.select_date(self.DataStart_entry)])
        self.CalendarOpenButton.grid(row=2, column=3, pady=5, padx=10, sticky="ew")
        
        
        customtkinter.CTkLabel(master=self.frame1,text="Выберите дату конца").grid(row=3, column=0, padx=10,columnspan=4, pady=10, sticky="ew")
        self.DataEnd_entry = customtkinter.CTkEntry(master=self.frame1, placeholder_text="ГГГГ-ММ-ДД")
        self.DataEnd_entry.grid(row=4, column=1,columnspan=2, padx=10, pady=10, sticky="ew")
        
        self.image_icon_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "calendarICO.png")), size=(20, 20))
        self.CalendarOpenButton = customtkinter.CTkButton(master=self.frame1, text="Выбрать дату",image=self.image_icon_image, command=lambda: [self.select_date(self.DataEnd_entry)])
        self.CalendarOpenButton.grid(row=4, column=3, pady=5, padx=10, sticky="ew")
        
        
        self.MakeReport1Button = customtkinter.CTkButton(master=self.frame1, text="Сформировать и открыть отчет", command=lambda: self.MakeReport1())
        self.MakeReport1Button.grid(row=5, column=0,columnspan=4, pady=5, padx=10, sticky="ew")
    def create_frame2(self):
        self.frame2 = customtkinter.CTkFrame(self,height=520)
        self.frame2.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
        self.frame2.grid_columnconfigure((0,1,2,3),weight=1)
        customtkinter.CTkLabel(master=self.frame2,text="Отчет по месту установки",fg_color="gray30", font=("Arial", 14)).grid(row=0,columnspan=4, column=0, padx=10, pady=10, sticky="ew")
        
        self.combobox1= customtkinter.CTkComboBox(master=self.frame2,values=[" "], state="readonly")
        self.combobox1.grid(row=1, column=0, padx=10, pady=10, sticky="ew")
        self.FillComboBoxes()
        
        self.MakeReport2Button = customtkinter.CTkButton(master=self.frame2, text="Сформировать и открыть отчет", command=lambda: self.MakeReport2())
        self.MakeReport2Button.grid(row=2, column=0,columnspan=4, pady=5, padx=10, sticky="ew")

        
    
        
        
    def create_frame3(self):
        self.frame3 = customtkinter.CTkScrollableFrame(self, height=920, width=450)
        self.grid_columnconfigure(2,weight=1)
        self.grid_rowconfigure((0,1),weight=1)
        self.frame3.grid(row=0,column=2,rowspan=3, padx=10, pady=10, sticky="ew")
        self.frame3.grid_columnconfigure((0,1,2), weight=1)
        customtkinter.CTkLabel(master=self.frame3, text="Отчет на выбранный компьютер", fg_color="gray30", font=("Arial", 14)).grid(row=0, columnspan=4, column=0, padx=10, pady=10, sticky="ew")
        self.combobox4= customtkinter.CTkComboBox(master=self.frame3,values=[" "], state="readonly")
        self.combobox4.grid(row=1, column=0, padx=10, pady=10, sticky="ew")
        self.FillComboBoxes4()
        
        self.MakeReport3Button = customtkinter.CTkButton(master=self.frame3, text="Сформировать и открыть отчет", command=lambda: self.MakeReport3())
        self.MakeReport3Button.grid(row=2, column=0,columnspan=4, pady=5, padx=10, sticky="ew")
    
    
    def create_frame4(self):

        customtkinter.CTkLabel(master=self.frame3, text="Отчет по ремонтам за промежуток дат", fg_color="gray30", font=("Arial", 14)).grid(row=3, columnspan=4,column=0, padx=10, pady=10, sticky="ew")
        
        
        customtkinter.CTkLabel(master=self.frame3,text="Выберите дату начала периода").grid(row=4, column=0, padx=10,columnspan=4, pady=10, sticky="ew")
        self.DataStart_entry1 = customtkinter.CTkEntry(master=self.frame3, placeholder_text="ГГГГ-ММ-ДД")
        self.DataStart_entry1.grid(row=5, column=0,columnspan=3, padx=10, pady=10, sticky="ew")
        
        self.image_icon_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "calendarICO.png")), size=(20, 20))
        self.CalendarOpenButton = customtkinter.CTkButton(master=self.frame3, text="Выбрать дату",image=self.image_icon_image, command=lambda: [self.select_date(self.DataStart_entry1)])
        self.CalendarOpenButton.grid(row=5, column=3, pady=5, padx=10, sticky="ew")
        
        
        customtkinter.CTkLabel(master=self.frame3,text="Выберите дату конца периода").grid(row=6, column=0, padx=10,columnspan=4, pady=10, sticky="ew")
        self.DataEnd_entry1 = customtkinter.CTkEntry(master=self.frame3, placeholder_text="ГГГГ-ММ-ДД")
        self.DataEnd_entry1.grid(row=7, column=0,columnspan=3, padx=10, pady=10, sticky="ew")
        
        self.image_icon_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "calendarICO.png")), size=(20, 20))
        self.CalendarOpenButton = customtkinter.CTkButton(master=self.frame3, text="Выбрать дату",image=self.image_icon_image, command=lambda: [self.select_date(self.DataEnd_entry1)])
        self.CalendarOpenButton.grid(row=7, column=3, pady=5, padx=10, sticky="ew")
        
        self.MakeReport4Button = customtkinter.CTkButton(master=self.frame3, text="Сформировать и открыть отчет", command=lambda: self.MakeReport4())
        self.MakeReport4Button.grid(row=8, column=0,columnspan=4, pady=5, padx=10, sticky="ew")
        
        
    def MakeReport4(self):
        date_start = self.DataStart_entry1.get()
        date_end = self.DataEnd_entry1.get()
        data=sql.get_repairs_during_period(date_start,date_end)
        # Создаем новый документ
        doc = Document()

        # Добавляем заголовок
        doc.add_heading(f'Отчет по ремонтам c {date_start} по {date_end}', level=1)

        # Создаем таблицу с заголовками
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = 'IP'
        hdr_cells[1].text = 'Имя ПК'
        hdr_cells[2].text = 'Описание ремонта'
        hdr_cells[3].text = 'Дата'

        # Добавляем данные в таблицу
        for item in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item[1])
            row_cells[1].text = item[2]
            row_cells[2].text = item[3]
            row_cells[3].text = item[4]

        # Сохраняем документ
        file_path = 'RepairsBDates.docx'
        # Удаление файла, если он существует
        if os.path.exists(file_path):
            os.remove(file_path)
        # Сохранение документа
        doc.save(file_path)
        # Открытие файла с помощью программы по умолчанию
        os.startfile(file_path)
        
        
        
    def select_date(self, entry):
        try:
            if hasattr(self, 'additionalWIN') and self.additionalWIN.winfo_exists():
                self.additionalWIN.focus()
                return
            self.additionalWIN = customtkinter.CTkToplevel(self)
            self.additionalWIN.geometry("260x200")
            self.additionalWIN.focus()
            self.additionalWIN.title(f"Calendar")
            self.additionalWIN.focus()
            maximize_minimize_button.hide(self.additionalWIN)
            cal = Calendar(self.additionalWIN, selectmode='day', date_pattern="yyyy-mm-dd")

            def set_date(selected_entry):
                def inner():
                    selected_date = cal.get_date()
                    selected_entry.delete(0, 'end')
                    selected_entry.insert(0, selected_date)
                    self.additionalWIN.destroy()
                return inner

            cal.pack()
            cal.bind('<<CalendarSelected>>', lambda event, entry=entry: set_date(entry)())
        except Exception as e:
            print(f"Exception in SelecTData: {e}")
    def MakeReport3(self):
        doc = DocxTemplate("PATTERN_REPORT_FOR_ONE.docx")
        currVal = self.combobox4.get()
        parts = currVal.split('|')  # Разделить строку на подстроки по символу '|'
        if len(parts) > 0:
            numbers = parts[0].strip()  # Получить первую подстроку и удалить лишние пробелы
        else:
            pass
        basic_data=sql.get_basic_info_by_id(numbers)
        detail_data=sql.get_detail_info_by_id(basic_data[0])
        component=sql.get_components_by_id(detail_data[2])
        
        _inv_numb=detail_data[3]
        _network_name=basic_data[2]
        _location=basic_data[3]
        _description_basic=basic_data[4]
        _last_status=basic_data[5]
        _data_laststatus=basic_data[6]
        _last_repair=basic_data[7]
        _serial_numb=detail_data[4]
        _mac_adr=detail_data[5]
        _wake_on_lan=detail_data[6]
        _vnc_password=detail_data[7]
        _processor=component[1]
        _ram=component[2]
        _motherboard=component[3]
        _graphicCard=component[4]
        _psu=component[5]
        _networkCard=component[6]
        _cooler=component[7]
        _chasis=component[8]
        _hdd=component[9]
        _ssd=component[10]
        _monitor=component[11]
        _keyboard=component[12]
        _mouse=component[13]
        _audio=component[14]
        context = {
        'inv_numb': _inv_numb,
        'network_name': _network_name,
        'location': _location,
        'description_basic': _description_basic,
        'last_status': _last_status,
        'data_laststatus': _data_laststatus,
        'last_repair': _last_repair,
        'serial_numb': _serial_numb,
        'mac_adr': _mac_adr,
        'wake_on_lan': _wake_on_lan,
        'vnc_password':  _vnc_password,
        'processor': _processor,
        'ram': _ram,
        'motherboard': _motherboard,
        'graphicCard': _graphicCard,
        'psu': _psu,
        'networkCard': _networkCard,
        'cooler': _cooler,
        'chasis': _chasis,
        'hdd': _hdd,
        'ssd': _ssd,
        'monitor': _monitor,
        'keyboard': _keyboard,
        'mouse': _mouse,
        'audio': _audio
        }

        doc.render(context)
        doc.add_page_break()

        # Создаем абзац
        paragraph = doc.add_paragraph()

        # Добавляем текст и задаем стиль
        run = paragraph.add_run("Статусы данного ПК")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        # Выравнивание текста
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
        statuses=sql.get_statuses_by_basic_info_id(basic_data[0])
        
        # Добавляем таблицу с двумя колонками
        table = doc.add_table(rows=len(statuses), cols=2)
        # Заполняем таблицу данными из списка statuses
        for i, status in enumerate(statuses):
            status_value = 'Вкл' if status[2] else 'Выкл'
            table.cell(i, 0).text = str(status_value)  # Значение столбца 1
            table.cell(i, 1).text = str(status[3])  # Значение столбца 2
        
        paragraph = doc.add_paragraph()
        run=paragraph.add_run("Ремонты:")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        repairs_list=[]
        repairs_list=sql.get_repairs_by_basic_id(numbers)
        
        number_of_tuples = sum(len(sublist) for sublist in repairs_list)

        table2=doc.add_table(rows=number_of_tuples, cols=2)
        for sublist in repairs_list:
            for i, rep in enumerate(sublist):               
                table2.cell(i, 0).text = str(rep[2])
                table2.cell(i, 1).text = str(rep[3])
            
        file_path="ReportPc.docx"
        # Удаление файла, если он существует
        if os.path.exists(file_path):
            os.remove(file_path)
        # Сохранение документа
        doc.save(file_path)
        # Открытие файла с помощью программы по умолчанию
        os.startfile(file_path)    
    def FillComboBoxes4(self):
        ToComboBoxOne = sql.get_basic_info()
        sasha = [str(data[0]) + " | " + str(data[1]) for data in ToComboBoxOne]
        self.combobox4.configure(values=sasha)
        self.update()
    def select_date1(self):
        try:
            if hasattr(self, 'additionalWIN') and self.additionalWIN.winfo_exists():
                self.additionalWIN.focus()
                return
            self.additionalWIN = customtkinter.CTkToplevel(self)
            self.additionalWIN.geometry("260x200")
            self.additionalWIN.focus()
            self.additionalWIN.title(f"Calendar")
            self.additionalWIN.focus()
            maximize_minimize_button.hide(self.additionalWIN)
            cal = Calendar(self.additionalWIN, selectmode='day',date_pattern="yyyy-mm-dd")
            def set_date():
                selected_date = cal.get_date()
                self.DataStart_entry.delete(0, 'end')  
                self.DataStart_entry.insert(0, selected_date)  
                self.additionalWIN.destroy() 
            cal.pack() 
            cal.bind('<<CalendarSelected>>', lambda event: set_date())
        except Exception as e:
            print(f"Exception in SelecTData: {e}")            
    def select_date2(self):
        try:
            if hasattr(self, 'additionalWIN') and self.additionalWIN.winfo_exists():
                self.additionalWIN.focus()
                return
            self.additionalWIN = customtkinter.CTkToplevel(self)
            self.additionalWIN.geometry("260x200")
            self.additionalWIN.focus()
            self.additionalWIN.title(f"Calendar")
            self.additionalWIN.focus()
            maximize_minimize_button.hide(self.additionalWIN)
            cal = Calendar(self.additionalWIN, selectmode='day',date_pattern="yyyy-mm-dd")
            def set_date():
                selected_date = cal.get_date()
                self.DataEnd_entry.delete(0, 'end')  
                self.DataEnd_entry.insert(0, selected_date)  
                self.additionalWIN.destroy() 
            cal.pack() 
            cal.bind('<<CalendarSelected>>', lambda event: set_date())
        except Exception as e:
            print(f"Exception in SelecTData: {e}")        
    def FillComboBoxes(self):
        ToComboBoxOne = sql.get_basic_info()
        qwe={''} ##все кабинетики (вкладочки)
        for i in range(len(ToComboBoxOne)):
            qwe.add(ToComboBoxOne[i][3])

        qwe.discard('')
        qwe = sorted(qwe)
        sasha = [str(data) for data in qwe]
        self.combobox1.configure(values=sasha)
        self.update()
    def MakeReport1(self):
        date_start = self.DataStart_entry.get()
        date_end = self.DataEnd_entry.get()
        date_start_obj = datetime.strptime(date_start, '%Y-%m-%d')
        date_end_obj = datetime.strptime(date_end, '%Y-%m-%d')
        date_start_formatted = date_start_obj.strftime('%Y-%d-%m')
        date_end_formatted = date_end_obj.strftime('%Y-%d-%m')

        data = sql.call_PCStatusDuringPeriod_procedure(date_start_formatted, date_end_formatted)
        doc = Document()
        title = doc.add_heading('PC Status Report', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Добавление таблицы с данными
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Заголовки столбцов
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'IP'
        hdr_cells[1].text = 'Location'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Date'

        # Форматирование текста
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    font = run.font
                    font.name = 'Arial'  # Название шрифта
                    font.size = Pt(11)    # Размер шрифта
                    font.bold = True      # Жирный шрифт
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру

        # Добавление данных из списка в таблицу
        for ip, location, status, date in data:
            row_cells = table.add_row().cells
            row_cells[0].text = ip
            row_cells[1].text = location
            status_text = 'On' if status else 'Off'
            row_cells[2].text = status_text
            # Установка цвета текста в ячейке в зависимости от статуса
            if status:
                color = RGBColor(0, 128, 0)  # Зеленый цвет для "On"
            else:
                color = RGBColor(255, 0, 0)  # Красный цвет для "Off"
            for paragraph in row_cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color

            row_cells[3].text = str(date)
        file_path = 'PC_Status_Report.docx'
        # Удаление файла, если он существует
        if os.path.exists(file_path):
            os.remove(file_path)
        # Сохранение документа
        doc.save(file_path)
        # Открытие файла с помощью программы по умолчанию
        os.startfile(file_path)
    def MakeReport2(self):
        currVal = self.combobox1.get()
        data=sql.get_device_info_by_location(currVal)
        
        try:
            # Создание нового документа
            doc = Document()

            # Заголовок документа
            title = doc.add_heading('Device Location Report', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Добавление таблицы с данными
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Заголовки столбцов
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'IP'
            hdr_cells[1].text = 'Location'
            hdr_cells[2].text = 'Last Repair'
            hdr_cells[3].text = 'Last Status'

            # Форматирование текста
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0]
                        font = run.font
                        font.name = 'Arial'  # Название шрифта
                        font.size = Pt(11)    # Размер шрифта
                        font.bold = True      # Жирный шрифт
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру

            # Добавление данных из списка в таблицу
            for ip, location, last_repair, status in data:
                row_cells = table.add_row().cells
                row_cells[0].text = ip
                row_cells[1].text = location
                row_cells[2].text = str(last_repair)
                status_text = 'On' if status else 'Off'
                row_cells[3].text = status_text
                # Установка цвета текста в ячейке в зависимости от статуса
                if status:
                    color = RGBColor(0, 128, 0)  # Зеленый цвет для "On"
                else:
                    color = RGBColor(255, 0, 0)  # Красный цвет для "Off"
                for paragraph in row_cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = color

            file_path = 'Device_Location_Report.docx'
            # Удаление файла, если он существует
            if os.path.exists(file_path):
                os.remove(file_path)
            # Сохранение документа
            doc.save(file_path)
            # Открытие файла с помощью программы по умолчанию
            os.startfile(file_path)

        except Exception as e:
            print(f"Error creating report: {str(e)}")
        
        #отчет по все инфо на выбранный пк

        #отчет по ремонтам за промежуток дат
        
if __name__ == "__main__":
    root = tkinter.Tk()
    app = Reports(root)
    root.mainloop()